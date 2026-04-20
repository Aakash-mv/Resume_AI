"""
Text extraction from PDF, DOCX, and TXT files.
"""
import os
import re

def normalize_text(text: str) -> str:
    """Basic text normalization: whitespace, artifacts, and encoding."""
    # Remove PDF artifacts like (cid:127)
    text = re.sub(r"\(cid:\d+\)", " ", text)
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove control characters
    text = "".join(ch for ch in text if ch.isprintable() or ch in ('\n', '\r', '\t'))
    return text.strip()

def extract_text_from_file(filepath: str, suffix: str) -> str:
    """Extract plain text from a file based on its extension."""
    suffix = suffix.lower()

    text = ""
    if suffix == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

    elif suffix == ".pdf":
        text = _extract_pdf(filepath)

    elif suffix in (".doc", ".docx"):
        text = _extract_docx(filepath)

    return normalize_text(text)


def _extract_pdf(filepath: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def _extract_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also grab tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")
